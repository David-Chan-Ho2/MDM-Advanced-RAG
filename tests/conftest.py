"""
Shared fixtures for all tests.
Creates temporary files for each supported document format.
"""

import pytest
from pathlib import Path


@pytest.fixture
def tmp_txt(tmp_path: Path) -> Path:
    f = tmp_path / "sample.txt"
    f.write_text(
        "Product: Widget Pro\nPart Number: WP-1234\nWeight: 2.5 kg\nLength: 150 mm\n"
        "Operating Temperature: -20°C to 85°C\nVoltage: 24V DC\nIP Rating: IP67\n"
    )
    return f


@pytest.fixture
def tmp_docx(tmp_path: Path) -> Path:
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()
    doc.add_heading("Widget Pro Datasheet", level=1)
    doc.add_paragraph("Part Number: WP-1234")
    doc.add_paragraph("Weight: 2.5 kg")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Attribute"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "IP Rating"
    table.cell(1, 1).text = "IP67"

    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def tmp_xlsx(tmp_path: Path) -> Path:
    try:
        import openpyxl
    except ImportError:
        pytest.skip("openpyxl not installed")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Specs"
    ws.append(["Attribute", "Value", "Unit"])
    ws.append(["Length", 150, "mm"])
    ws.append(["Weight", 2.5, "kg"])
    ws.append(["Voltage", 24, "V"])

    path = tmp_path / "sample.xlsx"
    wb.save(str(path))
    return path


@pytest.fixture
def tmp_html(tmp_path: Path) -> Path:
    f = tmp_path / "sample.html"
    f.write_text("""
    <html>
    <head><title>Spec Sheet</title></head>
    <body>
      <nav>Nav content to ignore</nav>
      <h1>Widget Pro</h1>
      <p>Part Number: WP-1234</p>
      <table>
        <tr><th>Attribute</th><th>Value</th></tr>
        <tr><td>IP Rating</td><td>IP67</td></tr>
        <tr><td>Weight</td><td>2.5 kg</td></tr>
      </table>
    </body>
    </html>
    """)
    return f
