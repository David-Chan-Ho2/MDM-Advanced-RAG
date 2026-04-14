"""
Generate synthetic product specification test documents in all supported formats.
Produces 5 files in data/raw/ — one per parser (PDF, DOCX, Excel, HTML, TXT).
"""

from pathlib import Path

OUT_DIR = Path(__file__).parent.parent / "data" / "raw"
OUT_DIR.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# 1. TXT — pressure sensor
# ---------------------------------------------------------------------------

def make_txt():
    path = OUT_DIR / "PS-4400_pressure_sensor.txt"
    path.write_text("""\
PRODUCT SPECIFICATION SHEET
============================
Manufacturer: Nexaflow Instrumentation
Product Family: ProSense Pressure Sensors
Part Number: PS-4400
SKU: NF-PS4400-SS
GTIN: 00614141123452
Revision: C

PHYSICAL DIMENSIONS
-------------------
Length: 95 mm
Width: 42 mm
Height: 38 mm
Weight: 210 g

ELECTRICAL SPECIFICATIONS
--------------------------
Voltage Input: 10–30 V DC
Current Rating: 4–20 mA
Power Consumption: 0.6 W
IP Rating: IP67

MATERIALS & COMPLIANCE
-----------------------
Primary Material: 316L Stainless Steel
Finish: Electropolished
RoHS Compliant: Yes
REACH Compliant: Yes
Certifications: CE, UL Listed, ATEX Zone 2

PERFORMANCE METRICS
-------------------
Operating Temperature (min): -40 °C
Operating Temperature (max): 85 °C
Storage Temperature (min): -50 °C
Storage Temperature (max): 100 °C
Accuracy: ±0.25% FS
Pressure Rating: 400 bar

DESCRIPTION
-----------
The PS-4400 is a robust industrial pressure transmitter designed for use in
demanding process environments. The 316L stainless steel wetted parts and
IP67 housing make it suitable for aggressive media and wash-down conditions.
The 4–20 mA two-wire output is compatible with most PLCs and DCS systems.
""")
    print(f"  wrote {path.name}")


# ---------------------------------------------------------------------------
# 2. HTML — flow meter
# ---------------------------------------------------------------------------

def make_html():
    path = OUT_DIR / "FM-2200_flow_meter.html"
    path.write_text("""\
<!DOCTYPE html>
<html lang="en">
<head><meta charset="UTF-8"><title>FM-2200 Flow Meter — Product Specification</title></head>
<body>
<header><nav>Home | Products | Support</nav></header>

<h1>FM-2200 Electromagnetic Flow Meter</h1>
<p><strong>Manufacturer:</strong> Aquatec Systems Ltd</p>
<p><strong>Part Number:</strong> FM-2200-DN50</p>
<p><strong>SKU:</strong> AQ-FM2200-050</p>
<p><strong>Product Family:</strong> AquaFlow Series</p>
<p><strong>Revision:</strong> B2</p>

<h2>Physical Dimensions</h2>
<table border="1">
  <tr><th>Attribute</th><th>Value</th><th>Unit</th></tr>
  <tr><td>Length (face-to-face)</td><td>200</td><td>mm</td></tr>
  <tr><td>Flange Diameter</td><td>165</td><td>mm</td></tr>
  <tr><td>Weight</td><td>3.8</td><td>kg</td></tr>
</table>

<h2>Electrical Specifications</h2>
<table border="1">
  <tr><th>Attribute</th><th>Value</th><th>Unit</th></tr>
  <tr><td>Supply Voltage</td><td>85–265</td><td>V AC</td></tr>
  <tr><td>Power Consumption</td><td>15</td><td>W max</td></tr>
  <tr><td>Output Signal</td><td>4–20 mA / HART</td><td>—</td></tr>
  <tr><td>IP Rating</td><td>IP68</td><td>—</td></tr>
</table>

<h2>Materials &amp; Compliance</h2>
<ul>
  <li>Body: Carbon Steel (epoxy-coated)</li>
  <li>Liner: PTFE</li>
  <li>Electrodes: Hastelloy C-276</li>
  <li>RoHS Compliant: Yes</li>
  <li>Certifications: CE, ATEX II 2G, FM Approved</li>
</ul>

<h2>Performance</h2>
<table border="1">
  <tr><th>Attribute</th><th>Value</th><th>Unit</th></tr>
  <tr><td>Flow Rate (max)</td><td>850</td><td>m³/h</td></tr>
  <tr><td>Accuracy</td><td>±0.3%</td><td>of reading</td></tr>
  <tr><td>Operating Temp (min)</td><td>-10</td><td>°C</td></tr>
  <tr><td>Operating Temp (max)</td><td>120</td><td>°C</td></tr>
  <tr><td>Pressure Rating</td><td>16</td><td>bar</td></tr>
</table>

<footer>© 2024 Aquatec Systems Ltd. All rights reserved.</footer>
</body>
</html>
""")
    print(f"  wrote {path.name}")


# ---------------------------------------------------------------------------
# 3. DOCX — temperature transmitter
# ---------------------------------------------------------------------------

def make_docx():
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("TT-8800 Temperature Transmitter — Product Specification", 0)

    def kv(label, value):
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    kv("Manufacturer", "Thermex Industrial")
    kv("Part Number", "TT-8800-4W")
    kv("SKU", "TX-TT8800-4W")
    kv("GTIN", "00614141987654")
    kv("Product Family", "ThermaLink Series")
    kv("Revision", "A3")

    doc.add_heading("Physical Dimensions", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Attribute"
    hdr[1].text = "Value"
    hdr[2].text = "Unit"
    dims = [
        ("Length", "120", "mm"),
        ("Diameter", "28", "mm"),
        ("Weight", "95", "g"),
    ]
    for attr, val, unit in dims:
        row = table.add_row().cells
        row[0].text = attr
        row[1].text = val
        row[2].text = unit

    doc.add_heading("Electrical Specifications", 1)
    kv("Supply Voltage", "10–36 V DC")
    kv("Output Signal", "4–20 mA / HART 7")
    kv("Power Consumption", "0.9 W")
    kv("IP Rating", "IP66 / IP67")

    doc.add_heading("Materials & Compliance", 1)
    kv("Housing", "Aluminium alloy (powder-coated)")
    kv("Wetted Parts", "316L Stainless Steel")
    kv("RoHS Compliant", "Yes")
    kv("REACH Compliant", "Yes")
    kv("Certifications", "CE, IECEx, ATEX Zone 1/21, SIL 2")

    doc.add_heading("Performance Metrics", 1)
    kv("Operating Temperature (min)", "-50 °C")
    kv("Operating Temperature (max)", "120 °C")
    kv("Storage Temperature (min)", "-55 °C")
    kv("Storage Temperature (max)", "130 °C")
    kv("Accuracy", "±0.1 °C")
    kv("Response Time (T63)", "< 3 s in water")

    path = OUT_DIR / "TT-8800_temperature_transmitter.docx"
    doc.save(path)
    print(f"  wrote {path.name}")


# ---------------------------------------------------------------------------
# 4. Excel — level sensor (multi-sheet)
# ---------------------------------------------------------------------------

def make_excel():
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()

    # --- Sheet 1: Overview ---
    ws1 = wb.active
    ws1.title = "Overview"
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="2E5FA3")

    def set_header(ws, row, col, text):
        cell = ws.cell(row=row, column=col, value=text)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    overview = [
        ("Field", "Value"),
        ("Manufacturer", "LevelTech GmbH"),
        ("Part Number", "LS-6600-R"),
        ("SKU", "LT-LS6600R"),
        ("GTIN", "04006381333445"),
        ("Product Family", "LevelPro Radar Series"),
        ("Revision", "D"),
    ]
    for r, (k, v) in enumerate(overview, start=1):
        set_header(ws1, r, 1, k) if r == 1 else ws1.cell(row=r, column=1, value=k)
        ws1.cell(row=r, column=2, value=v)
    ws1.column_dimensions["A"].width = 22
    ws1.column_dimensions["B"].width = 30

    # --- Sheet 2: Specifications ---
    ws2 = wb.create_sheet("Specifications")
    specs = [
        ("Category", "Attribute", "Value", "Unit"),
        ("Physical", "Length", 180, "mm"),
        ("Physical", "Diameter", 54, "mm"),
        ("Physical", "Weight", 680, "g"),
        ("Electrical", "Supply Voltage", "18–36 V DC", ""),
        ("Electrical", "Power Consumption", 1.2, "W"),
        ("Electrical", "Output", "4–20 mA HART", ""),
        ("Electrical", "IP Rating", "IP68", ""),
        ("Performance", "Operating Temp Min", -40, "°C"),
        ("Performance", "Operating Temp Max", 80, "°C"),
        ("Performance", "Storage Temp Min", -50, "°C"),
        ("Performance", "Storage Temp Max", 90, "°C"),
        ("Performance", "Accuracy", "±3 mm", ""),
        ("Performance", "Max Range", 30, "m"),
        ("Performance", "Pressure Rating", 40, "bar"),
    ]
    for c in range(1, 5):
        set_header(ws2, 1, c, specs[0][c - 1])
    for r, row in enumerate(specs[1:], start=2):
        for c, val in enumerate(row, start=1):
            ws2.cell(row=r, column=c, value=val)
    for c in range(1, 5):
        ws2.column_dimensions[get_column_letter(c)].width = 22

    # --- Sheet 3: Compliance ---
    ws3 = wb.create_sheet("Compliance")
    compliance = [
        ("Standard", "Status"),
        ("RoHS", "Compliant"),
        ("REACH", "Compliant"),
        ("CE Marking", "Yes"),
        ("ATEX II 1G Ex ia IIC T6 Ga", "Certified"),
        ("IECEx", "Certified"),
        ("SIL 2 (IEC 61508)", "Certified"),
    ]
    for c in range(1, 3):
        set_header(ws3, 1, c, compliance[0][c - 1])
    for r, (std, status) in enumerate(compliance[1:], start=2):
        ws3.cell(row=r, column=1, value=std)
        ws3.cell(row=r, column=2, value=status)
    ws3.column_dimensions["A"].width = 35
    ws3.column_dimensions["B"].width = 15

    path = OUT_DIR / "LS-6600_level_sensor.xlsx"
    wb.save(path)
    print(f"  wrote {path.name}")


# ---------------------------------------------------------------------------
# 5. PDF — motor drive (text-based, no images)
# ---------------------------------------------------------------------------

def make_pdf():
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "VFD-3300 Variable Frequency Drive", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 6, "Product Specification Sheet  |  Rev E", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    def section(title):
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_fill_color(46, 95, 163)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(0, 8, f"  {title}", fill=True, new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", "", 10)
        pdf.ln(2)

    def row(label, value):
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(65, 6, label + ":", border="B")
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 6, value, border="B", new_x="LMARGIN", new_y="NEXT")

    section("Product Identification")
    row("Manufacturer", "DriveTech Solutions")
    row("Part Number", "VFD-3300-7K5")
    row("SKU", "DT-VFD3300-7K5")
    row("GTIN", "00614141765432")
    row("Product Family", "VectorDrive 3000 Series")
    row("Revision", "E")
    pdf.ln(4)

    section("Physical Dimensions")
    row("Height", "320 mm")
    row("Width", "140 mm")
    row("Depth", "195 mm")
    row("Weight", "5.2 kg")
    pdf.ln(4)

    section("Electrical Specifications")
    row("Input Voltage", "380-480 V AC, 3-phase")
    row("Output Voltage", "0-480 V AC (variable)")
    row("Current Rating", "18 A")
    row("Power (motor)", "7.5 kW / 10 HP")
    row("Frequency", "50/60 Hz")
    row("IP Rating", "IP20 (IP55 optional)")
    pdf.ln(4)

    section("Materials & Compliance")
    row("Enclosure", "Powder-coated steel")
    row("RoHS Compliant", "Yes")
    row("REACH Compliant", "Yes")
    row("Certifications", "CE, UL 508C, cUL, RCM")
    pdf.ln(4)

    section("Performance Metrics")
    row("Operating Temp (min)", "-10 °C")
    row("Operating Temp (max)", "50 °C")
    row("Storage Temp (min)", "-40 °C")
    row("Storage Temp (max)", "70 °C")
    row("Accuracy (speed)", "±0.5% of set speed")
    row("Response Time", "< 50 ms torque step")
    pdf.ln(4)

    section("Description")
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(0, 6,
        "The VFD-3300 is a compact, high-performance variable frequency drive for "
        "three-phase AC induction and permanent magnet motors. Vector control with "
        "optional encoder feedback delivers precise speed and torque regulation "
        "across the full speed range. Built-in EMC filter (C2), dynamic braking "
        "transistor, and STO (Safe Torque Off, SIL 2) are standard. Suitable for "
        "pumps, fans, conveyors, and general-purpose drives."
    )

    path = OUT_DIR / "VFD-3300_variable_frequency_drive.pdf"
    pdf.output(str(path))
    print(f"  wrote {path.name}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print(f"Generating test documents in {OUT_DIR}/\n")
    make_txt()
    make_html()
    make_docx()
    make_excel()
    make_pdf()
    print("\nDone. 5 test documents created.")
