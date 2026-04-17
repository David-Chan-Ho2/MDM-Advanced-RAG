from pathlib import Path

from .base import BaseParser, ParsedDocument, ParsedPage


def _table_to_markdown(table) -> str:
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(rows)


class DocxParser(BaseParser):
    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in {".docx", ".doc"}

    def parse(self, file_path: Path) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("python-docx is required: pip install python-docx")

        doc = Document(str(file_path))

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)

        tables = [_table_to_markdown(t) for t in doc.tables]
        tables = [t for t in tables if t]

        return ParsedDocument(
            document_id=file_path.stem,
            filename=file_path.name,
            file_format="docx",
            pages=[ParsedPage(page_number=1, text=text, tables=tables)],
            metadata={"paragraph_count": len(paragraphs), "table_count": len(tables)},
        )
